"""AOF138 — rendu ÉDITABLE de la checklist partenaire (.docx).

Pourquoi un .docx et pas un PDF
================================
La checklist n'est pas une pièce à lire : c'est une pièce à REMPLIR par le
co-traitant — blancs du CPS, RIB, cases à cocher, mention manuscrite « lu et
accepté ». Un PDF y est inutilisable : le partenaire l'imprime, écrit dessus,
scanne, et l'ERP ne sait plus rien de l'état réel. Le format Word est ici la
fonction, pas un confort.

Dépendance OPTIONNELLE, jamais obligatoire
------------------------------------------
``python-docx`` (MIT, pur Python, sans bibliothèque système) est importé
**fonction-locale** et son absence est un cas NORMAL, pas une panne : le rendu
dégrade proprement en PDF portant la mention « pièce à fournir », exactement
comme ``pyHanko`` (scellement PAdES) et ``statsmodels`` dans ce dépôt.

Tant que la ligne ``python-docx`` n'est pas inscrite dans
``backend/django_core/requirements.txt`` (décision fondateur — la tâche est
tagguée ``@blocked: nouvelle dépendance``), c'est la voie dégradée qui
s'exécute en production et en CI. Le jour où la ligne est ajoutée, aucune
autre modification n'est nécessaire : le même appel produit un .docx.

L'alternative écartée : écrire du WordprocessingML à la main via ``zipfile``
— un à deux jours de travail fragile pour éviter une dépendance gratuite.

Contrat d'entrée
----------------
``blocs`` : liste de dicts ``{'code', 'titre', 'lignes': [...]}`` où chaque
ligne est ``{'libelle', 'obligatoire', 'cochee', 'responsable', 'commentaire'}``
— c'est la forme des lignes d'état de la checklist (AOF136), consommée telle
quelle. Ce module ne DÉCIDE de rien : il met en page.

Étanchéité (ratchet AOF129 étendu au format DOCX)
-------------------------------------------------
La checklist part chez un partenaire : elle est client-facing. Aucun libellé ne
peut y porter un prix d'achat, un coût de revient, une marge ou un bénéfice —
le contrôle est fait à l'ENTRÉE (refus) et le test relit le texte du .docx
produit, pas seulement la structure d'entrée.
"""
from __future__ import annotations

__all__ = [
    'FORMAT_DOCX',
    'FORMAT_PDF_DEGRADE',
    'MENTION_A_FOURNIR',
    'MOTS_INTERDITS',
    'docx_disponible',
    'html_degrade',
    'rendre_checklist',
    'texte_du_docx',
]

FORMAT_DOCX = 'docx'
FORMAT_PDF_DEGRADE = 'pdf'

CASE_VIDE = '☐'   # ☐
CASE_COCHEE = '☑'  # ☑

#: Mêmes racines que le lexique bloquant de la sanitisation (AOF143). Répétées
#: ici volontairement : ce module doit pouvoir refuser SEUL, même appelé hors
#: du chemin de sanitisation.
MOTS_INTERDITS = (
    "prix d'achat", 'prix achat', 'coût de revient', 'cout de revient',
    'marge', 'bénéfice', 'benefice', 'maximum posable',
)


def docx_disponible():
    """Vrai si ``python-docx`` est installé. Import fonction-local."""
    try:
        import docx  # noqa: F401
    except Exception:
        return False
    return True


def _verifier_etancheite(blocs):
    fautes = []
    for bloc in blocs or []:
        for ligne in bloc.get('lignes') or []:
            texte = ' '.join(
                str(ligne.get(cle) or '')
                for cle in ('libelle', 'commentaire', 'responsable')
            ).lower()
            for mot in MOTS_INTERDITS:
                if mot in texte:
                    fautes.append('{} / {} : « {} »'.format(
                        bloc.get('code', ''), ligne.get('libelle', ''), mot))
    if fautes:
        raise ValueError(
            "Checklist partenaire (pièce remise à un tiers) : mots réservés "
            "au directeur détectés — {}.".format(' ; '.join(fautes)))


def _entete(document, identite, marche):
    document.add_heading('Checklist de dépôt — à remplir par le partenaire', 0)
    ligne = document.add_paragraph()
    ligne.add_run('Soumissionnaire : ').bold = True
    ligne.add_run(str((identite or {}).get('raison_sociale') or ''))
    ligne = document.add_paragraph()
    ligne.add_run('Marché : ').bold = True
    ligne.add_run('{} — {}'.format((marche or {}).get('reference') or '',
                                   (marche or {}).get('objet') or ''))
    ligne = document.add_paragraph()
    ligne.add_run('Date de remise des plis : ').bold = True
    ligne.add_run(str((marche or {}).get('date_remise_plis') or ''))


def _rendre_docx(blocs, identite, marche):
    """Construit le .docx éditable. Appelé seulement si la lib est là."""
    import io

    from docx import Document

    document = Document()
    _entete(document, identite, marche)

    for bloc in blocs:
        document.add_heading(str(bloc.get('titre') or bloc.get('code') or ''),
                             level=1)
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        entetes = table.rows[0].cells
        for index, libelle in enumerate(
                ('Fait', 'Point à vérifier', 'Responsable', 'Observation')):
            entetes[index].text = libelle
        for ligne in bloc.get('lignes') or []:
            cellules = table.add_row().cells
            cellules[0].text = (CASE_COCHEE if ligne.get('cochee')
                                else CASE_VIDE)
            libelle = str(ligne.get('libelle') or '')
            if ligne.get('obligatoire'):
                libelle = '{} (obligatoire)'.format(libelle)
            cellules[1].text = libelle
            # Blancs à remplir À LA MAIN : la cellule reste vide, pas préremplie
            # d'un « — » qui ferait croire à un point traité.
            cellules[2].text = str(ligne.get('responsable') or '')
            cellules[3].text = str(ligne.get('commentaire') or '')
        document.add_paragraph()

    document.add_paragraph(
        "Mention manuscrite exigée sur chaque page du CPS : « lu et accepté », "
        "suivie du paraphe.")
    tampon = io.BytesIO()
    document.save(tampon)
    return tampon.getvalue()


MENTION_A_FOURNIR = 'PIÈCE À FOURNIR — version éditable indisponible'


def html_degrade(blocs, identite=None, marche=None):
    """HTML du repli « pièce à fournir ».

    Volontairement construit ICI et non dans un gabarit : le repli n'est pas
    une pièce du pack, c'est un CONSTAT d'indisponibilité. Lui donner un
    gabarit propre inviterait à l'utiliser comme une pièce normale.
    """
    from html import escape

    identite = identite or {}
    marche = marche or {}
    morceaux = [
        '<!doctype html><html lang="fr"><head><meta charset="utf-8">',
        '<title>Checklist de dépôt</title><style>',
        '@page{size:A4;margin:18mm 16mm;}',
        'body{font-family:"DejaVu Sans",Arial,sans-serif;font-size:9.5pt;}',
        'table{width:100%;border-collapse:collapse;margin-top:2mm;}',
        'th,td{border:1px solid #b5b5b5;padding:1.4mm 2mm;text-align:left;}',
        'th{background:#eee;}',
        '.alerte{border:2px solid #a00;color:#a00;padding:2mm 3mm;',
        'font-weight:bold;margin-bottom:4mm;}',
        '</style></head><body>',
        '<div class="alerte">{}</div>'.format(escape(MENTION_A_FOURNIR)),
        '<h1>Checklist de dépôt — à remplir par le partenaire</h1>',
        '<p><strong>Soumissionnaire :</strong> {}</p>'.format(
            escape(str(identite.get('raison_sociale') or ''))),
        '<p><strong>Marché :</strong> {} — {}</p>'.format(
            escape(str(marche.get('reference') or '')),
            escape(str(marche.get('objet') or ''))),
    ]
    for bloc in blocs or []:
        morceaux.append('<h2>{}</h2>'.format(
            escape(str(bloc.get('titre') or bloc.get('code') or ''))))
        morceaux.append('<table><thead><tr><th>Fait</th>'
                        '<th>Point à vérifier</th><th>Responsable</th>'
                        '<th>Observation</th></tr></thead><tbody>')
        for ligne in bloc.get('lignes') or []:
            libelle = escape(str(ligne.get('libelle') or ''))
            if ligne.get('obligatoire'):
                libelle = '{} (obligatoire)'.format(libelle)
            morceaux.append(
                '<tr><td>{}</td><td>{}</td><td>{}</td><td>{}</td></tr>'.format(
                    CASE_COCHEE if ligne.get('cochee') else CASE_VIDE,
                    libelle,
                    escape(str(ligne.get('responsable') or '')),
                    escape(str(ligne.get('commentaire') or ''))))
        morceaux.append('</tbody></table>')
    morceaux.append('</body></html>')
    return ''.join(morceaux)


def _rendre_degrade(blocs, identite, marche, company):
    """Repli PDF « pièce à fournir » quand ``python-docx`` est absent.

    Le PDF n'est PAS présenté comme la pièce définitive : il porte la mention
    « pièce à fournir », de sorte que le contrôleur de cohérence la compte
    comme non produite plutôt que verte à tort.
    """
    from core.pdf import render_pdf

    return render_pdf(html=html_degrade(blocs, identite, marche),
                      company=company)


def rendre_checklist(blocs, *, identite=None, marche=None, company=None,
                     forcer_degrade=False):
    """Rend la checklist. Renvoie ``(contenu, format, a_fournir)``.

    * ``format`` vaut ``'docx'`` (éditable) ou ``'pdf'`` (dégradé) ;
    * ``a_fournir`` est vrai en mode dégradé — la pièce est alors marquée
      « à fournir » et ne peut pas être comptée comme produite.
    """
    blocs = list(blocs or [])
    _verifier_etancheite(blocs)
    if forcer_degrade or not docx_disponible():
        return (_rendre_degrade(blocs, identite, marche, company),
                FORMAT_PDF_DEGRADE, True)
    return _rendre_docx(blocs, identite, marche), FORMAT_DOCX, False


def texte_du_docx(contenu):
    """Extrait le texte d'un .docx (pour les contrôles d'étanchéité/tests).

    Lit le XML du paquet OOXML directement : la vérification d'étanchéité ne
    doit pas dépendre de la bibliothèque qu'elle contrôle.
    """
    import io
    import re
    import zipfile

    morceaux = []
    with zipfile.ZipFile(io.BytesIO(contenu)) as paquet:
        for nom in paquet.namelist():
            if nom.startswith('word/') and nom.endswith('.xml'):
                xml = paquet.read(nom).decode('utf-8', 'replace')
                # `<w:t ...>` STRICTEMENT : `<w:t[^>]*>` attraperait aussi
                # `<w:tblPr>` et ferait entrer du balisage dans le « texte »,
                # ce qui rendrait tout contrôle d'étanchéité illisible.
                morceaux.extend(re.findall(r'<w:t(?:\s[^>]*)?>(.*?)</w:t>',
                                           xml, flags=re.S))
    return '\n'.join(morceaux)
